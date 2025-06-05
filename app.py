from flask import Flask, render_template, request, jsonify, send_file
import requests
import os
from docx import Document
from fpdf import FPDF

app = Flask(__name__)
generated_resume = ""

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/generate", methods=["POST"])
def generate_resume():
    global generated_resume
    data = request.json

    prompt = f"""
    Create a {data['template']} style resume:
    Name: {data['name']}
    Address: {data['address']}
    Phone: {data['phone']}
    Email: {data['email']}
    Summary: {data['summary']}
    Education: {data['education']}
    Work Experience: {data['experience']}
    Skills: {data['skills']}
    References: {data['references']}
    """

    response = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {os.environ.get('OPENROUTER_API_KEY')}",
            "Content-Type": "application/json"
        },
        json={
            "model": "openai/gpt-3.5-turbo",
            "messages": [{"role": "user", "content": prompt}]
        }
    )

    generated_resume = response.json()["choices"][0]["message"]["content"]
    return jsonify({"output": generated_resume})

@app.route("/download/pdf")
def download_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True)
    pdf.set_font("Arial", size=12)
    for line in generated_resume.split('\n'):
        pdf.multi_cell(0, 10, line)
    pdf.output("resume.pdf")
    return send_file("resume.pdf", as_attachment=True)

@app.route("/download/docx")
def download_docx():
    doc = Document()
    for line in generated_resume.split('\n'):
        doc.add_paragraph(line)
    doc.save("resume.docx")
    return send_file("resume.docx", as_attachment=True)
    from flask import make_response
from docx import Document
from fpdf import FPDF

@app.route('/download-pdf', methods=['POST'])
def download_pdf():
    data = request.json
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    pdf.cell(200, 10, txt="Resume", ln=True, align='C')
    pdf.ln(10)
    for key, value in data.items():
        pdf.multi_cell(0, 10, f"{key.capitalize()}: {value}")
        pdf.ln(1)

    response = make_response(pdf.output(dest='S').encode('latin1'))
    response.headers.set('Content-Disposition', 'attachment', filename='resume.pdf')
    response.headers.set('Content-Type', 'application/pdf')
    return response

@app.route('/download-docx', methods=['POST'])
def download_docx():
    data = request.json
    doc = Document()
    doc.add_heading('Resume', 0)

    for key, value in data.items():
        doc.add_heading(key.capitalize(), level=1)
        doc.add_paragraph(value)

    response = make_response()
    doc.save("resume.docx")
    with open("resume.docx", "rb") as f:
        response.data = f.read()

    response.headers.set('Content-Disposition', 'attachment', filename='resume.docx')
    response.headers.set('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    return response


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=81)
